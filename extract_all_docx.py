import os
from docx import Document
import glob

docx_files = glob.glob("recepten/**/*.docx", recursive=True) + glob.glob("*.docx")

for f in docx_files:
    print(f"--- FILE: {f} ---")
    try:
        doc = Document(f)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        print(f"TABLE_CELL: {cell.text.strip()}")
    except Exception as e:
        print(f"Error reading {f}: {e}")
    print("\n")
