import os
from docx import Document
import glob
import json

def clean_text(text):
    if not text:
        return ""
    # Replace common issues
    text = text.replace("IngrediÙnten", "Ingrediënten")
    text = text.replace("òá", "- ")
    text = text.replace("½", "1/2")
    text = text.replace("¼", "1/4")
    text = text.replace("¢", "1/2") # Common misreading of 1/2 in some encodings
    return text.strip()

docx_files = glob.glob("recepten/**/*.docx", recursive=True) + glob.glob("*.docx")

all_data = {}

for f in docx_files:
    filename = os.path.basename(f)
    print(f"Processing {f}...")
    try:
        doc = Document(f)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(clean_text(para.text))
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(clean_text(cell.text))
                if row_data:
                    content.append(" | ".join(row_data))
        
        all_data[f] = "\n".join(content)
    except Exception as e:
        print(f"Error reading {f}: {e}")

with open("docx_content.json", "w", encoding="utf-8") as f:
    json.dump(all_data, f, ensure_ascii=False, indent=2)
